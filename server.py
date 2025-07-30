#server
from mcp.server.fastmcp import FastMCP
from pydantic import BaseModel, Field
from mcp.server.fastmcp.prompts import base
from typing import TypedDict
import arxiv
import json
import os
from typing import List
from typing import Union
import re
import PyPDF2
from io import BytesIO
import mcp.types as types
try:
    import docx
except ImportError:
    docx = None
import tempfile
import base64
from pathlib import Path
from mcp import ClientSession, StdioServerParameters
from mcp.server.models import InitializationOptions
import mcp.server.stdio
from mcp.server import NotificationOptions, Server
from mcp.server.models import InitializationOptions
import mcp
from pydantic import BaseModel
from typing import Any, Sequence, List, Union
import httpx
from langchain_huggingface import HuggingFaceEmbeddings
from pinecone import Pinecone
from tavily import TavilyClient
from dotenv import load_dotenv
load_dotenv()
# File generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from fpdf import FPDF
import markdown
from jinja2 import Template
import uuid
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


#constants
NWS_API_BASE = "https://api.weather.gov"
USER_AGENT = "weather-app/1.0"

#Helper functions
async def make_nws_request(url: str) -> dict[str, Any] | None:
    """Make a request to the NWS API with proper error handling."""
    headers = {
        "User-Agent": USER_AGENT,
        "Accept": "application/geo+json"
    }
    async with httpx.AsyncClient() as client:
        try:
            response = await client.get(url, headers=headers, timeout=30.0)
            response.raise_for_status()
            return response.json()
        except Exception:
            return None
        
def format_alert(feature: dict) -> str:
    """Format an alert feature into a readable string."""
    props = feature["properties"]
    return f"""
Event: {props.get('event', 'Unknown')}
Area: {props.get('areaDesc', 'Unknown')}
Severity: {props.get('severity', 'Unknown')}
Description: {props.get('description', 'No description available')}
Instructions: {props.get('instruction', 'No specific instructions provided')}
"""


#Connection to Pinecone
pinecone_api_key = os.environ.get('PINECONE_API_KEY')
pc = Pinecone(api_key=pinecone_api_key)
index = pc.Index("sample3")
embedder = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

def sanitize_filename(name: str) -> str:
    # Remove all non-alphanumeric, dash, underscore characters
    return re.sub(r'[^a-zA-Z0-9-_]', '_', name)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PAPER_DIR = os.path.join(BASE_DIR, "papers")
GENERATED_FILES_DIR = os.path.join(BASE_DIR, "generated_files")


#create an MCP server
mcp = FastMCP("Research-Demo")

#Add an addition tool
@mcp.tool()
def sum(a: int, b: int) -> int:
    """Add two numbers"""
    return a+b


@mcp.tool()
async def get_alerts(state: str) -> str:
    """Get weather alerts for a US state.

    Args:
        state: Two-letter US state code (e.g. CA, NY)
    """
    url = f"{NWS_API_BASE}/alerts/active/area/{state}"
    data = await make_nws_request(url)

    if not data or "features" not in data:
        return "Unable to fetch alerts or no alerts found."

    if not data["features"]:
        return "No active alerts for this state."

    alerts = [format_alert(feature) for feature in data["features"]]
    return "\n---\n".join(alerts)


@mcp.tool()
async def get_forecast(latitude: float, longitude: float) -> str:
    """Get weather forecast for a location.

    Args:
        latitude: Latitude of the location
        longitude: Longitude of the location
    """
    # First get the forecast grid endpoint
    points_url = f"{NWS_API_BASE}/points/{latitude},{longitude}"
    points_data = await make_nws_request(points_url)

    if not points_data:
        return "Unable to fetch forecast data for this location."

    # Get the forecast URL from the points response
    forecast_url = points_data["properties"]["forecast"]
    forecast_data = await make_nws_request(forecast_url)

    if not forecast_data:
        return "Unable to fetch detailed forecast."

    # Format the periods into a readable forecast
    periods = forecast_data["properties"]["periods"]
    forecasts = []
    for period in periods[:5]:  # Only show next 5 periods
        forecast = f"""
{period['name']}:
Temperature: {period['temperature']}°{period['temperatureUnit']}
Wind: {period['windSpeed']} {period['windDirection']}
Forecast: {period['detailedForecast']}
"""
        forecasts.append(forecast)

    return "\n---\n".join(forecasts)

@mcp.tool()
def search_papers(topic: str, max_results: int = 5) -> List[str]:
    """
    Search for papers on arXiv based on a topic and store their information.
    
    Args:
        topic: The topic to search for
        max_results: Maximum number of results to retrieve (default: 5)
        
    Returns:
        List of paper IDs found in the search
    """
    
    # Use arxiv to find the papers 

    # Search for the most relevant articles matching the queried topic
    search = arxiv.Search(
        query = topic,
        max_results = max_results,
        sort_by = arxiv.SortCriterion.Relevance
    )

    papers = list(search.results())
    
    # Create directory for this topic
    path = os.path.join(PAPER_DIR, sanitize_filename(topic.lower()))
    os.makedirs(path, exist_ok=True)
    
    file_path = os.path.join(path, "papers_info.json")

    # Try to load existing papers info
    try:
        with open(file_path, "r") as json_file:
            papers_info = json.load(json_file)
    except (FileNotFoundError, json.JSONDecodeError):
        papers_info = {}

    # Process each paper and add to papers_info  
    paper_ids = []
    for paper in papers:
        paper_ids.append(paper.get_short_id())
        paper_info = {
            'title': paper.title,
            'authors': [author.name for author in paper.authors],
            'summary': paper.summary,
            'pdf_url': paper.pdf_url,
            'published': str(paper.published.date())
        }
        papers_info[paper.get_short_id()] = paper_info
    
    # Save updated papers_info to json file
    with open(file_path, "w") as json_file:
        json.dump(papers_info, json_file, indent=2)
    
    print(f"Results are saved in: {file_path}")
    
    return paper_ids

@mcp.tool()
def extract_info(paper_id: str) -> Union[str, None]:
    """
    Search for information about a specific paper across all topic directories.
    
    Args:
        paper_id: The ID of the paper to look for
        
    Returns:
        JSON string with paper information if found, error message if not found
    """
 
    for item in os.listdir(PAPER_DIR):
        item_path = os.path.join(PAPER_DIR, item)
        if os.path.isdir(item_path):
            file_path = os.path.join(item_path, "papers_info.json")
            if os.path.isfile(file_path):
                try:
                    with open(file_path, "r") as json_file:
                        papers_info = json.load(json_file)
                        if paper_id in papers_info:
                            return json.dumps(papers_info[paper_id], indent=2)
                except (FileNotFoundError, json.JSONDecodeError) as e:
                    print(f"Error reading {file_path}: {str(e)}")
                    continue
    
    return f"There's no saved information related to paper {paper_id}."

@mcp.tool()
def Broadaxis_knowledge_search(query: str):

    """
    Retrieves the most relevant company's(Broadaxus) information from the internal knowledge base in response to an company related query.

    This tool performs semantic search over a RAG-powered database containing details about the Broadaxis's background, team, projects, responsibilities, and domain expertise. It is designed to support tasks such as retreiving the knowledge regarding the company, surfacing domain-specific experience.

    Args:
        query: A natural language request related to the company’s past work, expertise, or capabilities (e.g., "What are the team's responsibilities?").
    """
    # try:
        # Step 1: Embed the query
    query_embedding = embedder.embed_documents([query])[0]

        # Step 2: Perform similarity search using Pinecone
    query_result = index.query(
        vector=[query_embedding],
        top_k=5,  # Recommend using more than 1 to allow LLM flexibility
        include_metadata=True,
        namespace=""  # Set if needed
    )
    documents = [result['metadata']['text'] for result in query_result['matches']]

    return documents

    # except Exception as e:
    #     return json.dumps({"error": str(e)})

# Initialize client (use env variable or hardcode the API key if preferred)

os.environ["TAVILY_API_KEY"] = "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"
tavily = TavilyClient()  # or TavilyClient(api_key="your_key")

@mcp.tool()
def web_search_tool(query: str):
    """
    Performs a real-time web search using Tavily and returns relevant results
    (including title, URL, and snippet).

    Args:
        query: A natural language search query.

    Returns:
        A JSON string with the top search results.
    """
    try:
        # Perform the search
        results = tavily.search(query=query, search_depth="advanced", include_answer=False)

        # Extract and format results
        formatted = [
            {
                "title": r.get("title"),
                "url": r.get("url"),
                "snippet": r.get("content")
            }
            for r in results.get("results", [])
        ]

        return json.dumps({"results": formatted})

    except Exception as e:
        return json.dumps({"error": str(e)})
    
@mcp.tool()
def generate_pdf_document(title: str, content: str, filename: str = None) -> str:
    """
    Generate a PDF document with the provided title and content.

    Args:
        title: The title of the document
        content: The main content of the document (supports markdown formatting)
        filename: Optional custom filename (without extension)

    Returns:
        JSON string with file information including download path
    """
    try:
        # Generate unique filename if not provided
        if not filename:
            filename = f"document_{uuid.uuid4().hex[:8]}"

        # Ensure filename doesn't have extension
        filename = filename.replace('.pdf', '')

        # Create full file path
        file_path = os.path.join(GENERATED_FILES_DIR, f"{filename}.pdf")

        # Create PDF document
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))

        # Process content (convert markdown to HTML-like formatting for reportlab)
        content_lines = content.split('\n')
        for line in content_lines:
            if line.strip():
                # Handle basic markdown formatting
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                elif line.startswith('- ') or line.startswith('* '):
                    story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
                else:
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))

        # Build PDF
        doc.build(story)

        # Get file info
        file_size = os.path.getsize(file_path)

        return json.dumps({
            "status": "success",
            "filename": f"{filename}.pdf",
            "file_path": file_path,
            "file_size": file_size,
            "download_url": f"/download/{filename}.pdf",
            "created_at": datetime.datetime.now().isoformat(),
            "type": "pdf"
        })

    except Exception as e:
        return json.dumps({
            "status": "error",
            "error": str(e)
        })


@mcp.tool()
def generate_word_document(title: str, content: str, filename: str = None) -> str:
    """
    Generate a Word document with the provided title and content.

    Args:
        title: The title of the document
        content: The main content of the document (supports basic markdown formatting)
        filename: Optional custom filename (without extension)

    Returns:
        JSON string with file information including download path
    """
    try:
        # Generate unique filename if not provided
        if not filename:
            filename = f"document_{uuid.uuid4().hex[:8]}"

        # Ensure filename doesn't have extension
        filename = filename.replace('.docx', '').replace('.doc', '')

        # Create full file path
        file_path = os.path.join(GENERATED_FILES_DIR, f"{filename}.docx")

        # Create Word document
        doc = Document()

        # Add title
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add some space after title
        doc.add_paragraph()

        # Process content (handle basic markdown formatting)
        content_lines = content.split('\n')
        for line in content_lines:
            line = line.strip()
            if line:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    # Add bullet point
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    # Add numbered list
                    doc.add_paragraph(line[3:], style='List Number')
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
            else:
                # Add empty paragraph for spacing
                doc.add_paragraph()

        # Save document
        doc.save(file_path)

        # Get file info
        file_size = os.path.getsize(file_path)

        return json.dumps({
            "status": "success",
            "filename": f"{filename}.docx",
            "file_path": file_path,
            "file_size": file_size,
            "download_url": f"/download/{filename}.docx",
            "created_at": datetime.datetime.now().isoformat(),
            "type": "docx"
        })

    except Exception as e:
        return json.dumps({
            "status": "error",
            "error": str(e)
        })


@mcp.tool()
def generate_text_file(content: str, filename: str = None, file_extension: str = "txt") -> str:
    """
    Generate a text file with the provided content.

    Args:
        content: The content to write to the file
        filename: Optional custom filename (without extension)
        file_extension: File extension (txt, md, csv, json, etc.)

    Returns:
        JSON string with file information including download path
    """
    try:
        # Generate unique filename if not provided
        if not filename:
            filename = f"file_{uuid.uuid4().hex[:8]}"

        # Clean filename and ensure no extension
        filename = filename.split('.')[0]

        # Create full file path
        file_path = os.path.join(GENERATED_FILES_DIR, f"{filename}.{file_extension}")

        # Write content to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

        # Get file info
        file_size = os.path.getsize(file_path)

        return json.dumps({
            "status": "success",
            "filename": f"{filename}.{file_extension}",
            "file_path": file_path,
            "file_size": file_size,
            "download_url": f"/download/{filename}.{file_extension}",
            "created_at": datetime.datetime.now().isoformat(),
            "type": file_extension
        })

    except Exception as e:
        return json.dumps({
            "status": "error",
            "error": str(e)
        })


@mcp.tool()
def list_generated_files() -> str:
    """
    List all generated files available for download.

    Returns:
        JSON string with list of available files
    """
    try:
        files = []
        if os.path.exists(GENERATED_FILES_DIR):
            for filename in os.listdir(GENERATED_FILES_DIR):
                file_path = os.path.join(GENERATED_FILES_DIR, filename)
                if os.path.isfile(file_path):
                    file_size = os.path.getsize(file_path)
                    file_modified = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))

                    files.append({
                        "filename": filename,
                        "file_size": file_size,
                        "download_url": f"/download/{filename}",
                        "modified_at": file_modified.isoformat(),
                        "type": filename.split('.')[-1] if '.' in filename else "unknown"
                    })

        return json.dumps({
            "status": "success",
            "files": files,
            "count": len(files)
        })

    except Exception as e:
        return json.dumps({
            "status": "error",
            "error": str(e)
        })


@mcp.tool()
def cleanup_old_files(days_old: int = 0) -> str:
    """
    Clean up generated files older than specified days.

    Args:
        days_old: Number of days old files should be to be deleted (default: 7)

    Returns:
        JSON string with cleanup results
    """
    try:
        deleted_files = []
        current_time = datetime.datetime.now()
        cutoff_time = current_time - datetime.timedelta(days=days_old)

        if os.path.exists(GENERATED_FILES_DIR):
            for filename in os.listdir(GENERATED_FILES_DIR):
                file_path = os.path.join(GENERATED_FILES_DIR, filename)
                if os.path.isfile(file_path):
                    file_modified = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))
                    if file_modified < cutoff_time:
                        os.remove(file_path)
                        deleted_files.append({
                            "filename": filename,
                            "modified_at": file_modified.isoformat()
                        })

        return json.dumps({
            "status": "success",
            "deleted_files": deleted_files,
            "count": len(deleted_files),
            "cutoff_date": cutoff_time.isoformat()
        })

    except Exception as e:
        return json.dumps({
            "status": "error",
            "error": str(e)
        })


class FileContent(BaseModel):
    name: str
    content: str



class FileContent(BaseModel):
    name: str
    content: str

@mcp.resource("file://desktop/{name}")
def read_document(name: str) -> FileContent:
    """Read a document and return its name and content"""
    safe_name = os.path.basename(name)
    file_path = os.path.join("C:/Users/rohka/OneDrive/Desktop", safe_name)
    
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    return FileContent(name=safe_name, content=content)


#add a dynamic greeting resource
@mcp.resource("greeting://{name}")
def get_greeting(name: str) -> str:
    """Get a personalized greeting"""
    return f"Hello, {name}!"

@mcp.resource("config://settings")
def get_settings() -> str:
    """Get application settings."""
    return """{
    "theme": "dark",
    "language": "en",
    "debug": false
    }"""

@mcp.resource("config://word_formatting_instructions")
def get_word_formatting_instructions() -> str:
    """Get Word document formatting instructions."""
    return """{
    "font": "Calibri",
    "font_size": 11,
    "heading_font_size": 14,
    "heading_bold": true,
    "line_spacing": 1.15,
    "margins": {
        "top": "1 inch",
        "bottom": "1 inch",
        "left": "1 inch",
        "right": "1 inch"
    },
    "page_orientation": "portrait"
    }"""


@mcp.prompt(title="Code Review")
def review_code(code: str) -> str:
    """Review a piece of code"""
    return f"Please review this code:\n\n{code}"


@mcp.prompt(title="Debug Assistant")
def debug_error(error: str) -> list[base.Message]:
    """Debug an error message"""
    return [
        base.UserMessage("I'm seeing this error:"),
        base.UserMessage(error),
        base.AssistantMessage("I'll help debug that. What have you tried so far?"),
    ]


@mcp.prompt(title="Movie Actor Analysis")
def analyze_actor(actor_name: str) -> str:
    """Analyze a movie actor's career, top movies, and performance"""
    return f"""Please provide a comprehensive analysis of the actor {actor_name}. Include the following information:

1. **Actor Overview:**
   - Brief biography and career highlights
   - Years active in the industry
   - Notable achievements and awards

2. **Top 5 Highest Grossing Movies:**
   - List their top 5 highest grossing films with box office numbers
   - Brief description of their role in each movie
   - Year of release for each film

3. **Genre Analysis:**
   - Primary genres they work in (Action, Drama, Comedy, etc.)
   - Which genres they excel in most
   - Any genre transitions throughout their career

4. **Performance Strengths:**
   - What makes them a compelling actor
   - Signature acting style or techniques
   - Most memorable performances

5. **Areas for Improvement:**
   - What could they have done better in their career
   - Missed opportunities or poor role choices
   - Suggestions for future career directions

6. **Overall Assessment:**
   - Current status in Hollywood
   - Legacy and impact on cinema
   - Comparison with peers in their generation

Please provide detailed, factual information with specific examples and be objective in your analysis."""

if __name__ == "__main__":
    # Initialize and run the server
    import sys
    import logging

    # Set up logging for debugging
    logging.basicConfig(level=logging.DEBUG)

    try:
        mcp.run(transport='stdio')
    except Exception as e:
        print(f"Server error: {e}", file=sys.stderr)
        sys.exit(1)